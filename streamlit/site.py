### IMPORTS ###
import streamlit as st
import zipfile
import io
from io import BytesIO
import pandas as pd
from docx import Document
from time import sleep

import os
import base64

import shutil
from PIL import Image
import imagehash

### FUNCTION DEFINITIONS ###


# return docx_files, placeholder_logo, localisation_csv
def extract_package_contents(package_zip):
    # Extract all .docx files ignoring system files
    docx_files = [
        f.split("/")[-1]
        for f in package_zip.namelist()
        if f.endswith(".docx") and "__MACOSX" not in f
    ]

    image_extensions = [".png", ".jpg", ".jpeg"]
    image_files = [
        f
        for f in package_zip.namelist()
        if any(f.lower().endswith(ext) for ext in image_extensions)
        and "__MACOSX" not in f
    ]
    if not image_files:
        placeholder_logo = None
    else:
        placeholder_logo = image_files[0]

    # Extract 'localisation.csv' files
    csv_files = [f for f in package_zip.namelist() if "localisation.csv" in f]
    if not csv_files:
        localisation_csv = None
    else:
        localisation_csv = csv_files[0]

    return docx_files, placeholder_logo, localisation_csv


def replace_text_in_paragraphs(paragraphs, changes):
    for paragraph in paragraphs:
        full_text = "".join(run.text for run in paragraph.runs)
        for old_text, new_text in changes.items():
            if old_text in full_text:
                # Replace text and rebuild paragraph
                new_text = full_text.replace(old_text, new_text)
                # Clear all runs
                for run in paragraph.runs:
                    run.text = ""
                # Add new run with updated text
                paragraph.add_run(new_text)
                break  # exit after first match to avoid duplicating replacements


def replace_text_in_tables(tables, changes):
    for table in tables:
        for row in table.rows:
            for cell in row.cells:
                replace_text_in_paragraphs(cell.paragraphs, changes)


def replace_text_in_headers_footers(sections, changes):
    for section in sections:
        # Check and replace text in default header and footer
        replace_text_in_paragraphs(section.header.paragraphs, changes)
        replace_text_in_paragraphs(section.footer.paragraphs, changes)
        # Check and replace text in first page header and footer if they exist
        if section.first_page_header is not None:
            replace_text_in_paragraphs(section.first_page_header.paragraphs, changes)
        if section.first_page_footer is not None:
            replace_text_in_paragraphs(section.first_page_footer.paragraphs, changes)
        # Check and replace text in even page header and footer if they exist
        if section.even_page_header is not None:
            replace_text_in_paragraphs(section.even_page_header.paragraphs, changes)
        if section.even_page_footer is not None:
            replace_text_in_paragraphs(section.even_page_footer.paragraphs, changes)


# # Function to extract images from a docx file
def extract_images(doc_bytes, extract_dir):
    """Unzip a docx file from bytes and extract all images to a specified directory."""
    with zipfile.ZipFile(io.BytesIO(doc_bytes), "r") as zip_ref:
        zip_ref.extractall(extract_dir)
    return [
        os.path.join(extract_dir, "word/media", name)
        for name in os.listdir(os.path.join(extract_dir, "word/media"))
    ]


# Function to find the matching image based on hash comparison
def find_matching_image(images, placeholder_image_bytes):
    """Find and return the path of the image that matches the placeholder image."""
    placeholder_hash = imagehash.average_hash(
        Image.open(io.BytesIO(placeholder_image_bytes))
    )
    for image_path in images:
        current_hash = imagehash.average_hash(Image.open(image_path))
        if current_hash - placeholder_hash == 0:  # comparing hashes
            return image_path
    return None


# Function to replace a specified image in the docx file with a new image
def replace_image(doc_bytes, placeholder_image_bytes, new_image_bytes):
    """Replace a specified image in the docx file with a new image."""
    temp_dir = "temp_docx"
    if os.path.exists(temp_dir):
        shutil.rmtree(temp_dir)
    os.makedirs(temp_dir)

    with zipfile.ZipFile(io.BytesIO(doc_bytes), "r") as zip_ref:
        zip_ref.extractall(temp_dir)

    media_dir = os.path.join(temp_dir, "word", "media")
    if not os.path.exists(media_dir):
        raise FileNotFoundError(f"{media_dir} does not exist")

    placeholder_hash = imagehash.average_hash(
        Image.open(io.BytesIO(placeholder_image_bytes))
    )

    replaced = False
    for image_name in os.listdir(media_dir):
        image_path = os.path.join(media_dir, image_name)
        with open(image_path, "rb") as img_file:
            img = Image.open(img_file)
            img_hash = imagehash.average_hash(img)
            if img_hash == placeholder_hash:

                # Check the new image bytes before writing
                if len(new_image_bytes) == 0:
                    break

                try:
                    # Attempt to open new image to verify it's valid
                    new_img = Image.open(io.BytesIO(new_image_bytes))
                except Exception as e:
                    break

                with open(image_path, "wb") as new_img_file:
                    new_img_file.write(new_image_bytes)
                replaced = True
                break

    if not replaced:
        st.write("Placeholder image not found in the document.")

    new_docx_io = io.BytesIO()
    with zipfile.ZipFile(new_docx_io, "w") as docx_zip:
        for foldername, subfolders, filenames in os.walk(temp_dir):
            for filename in filenames:
                file_path = os.path.join(foldername, filename)
                arcname = os.path.relpath(file_path, temp_dir)
                docx_zip.write(file_path, arcname)
    shutil.rmtree(temp_dir)
    # st.write("Temporary files cleaned up.")

    new_docx_io.seek(0)
    # st.write("Returning new docx bytes.")
    return new_docx_io.getvalue()


# Function to apply changes to documents
def apply_changes_to_document(
    changes_json, docx_files, placeholder_logo_bytes, new_logo_bytes
):
    modified_docs = {}
    for doc_name, templates in changes_json.items():
        doc = docx_files[doc_name]

        # Save document temporarily to replace the logo
        temp_doc_io = BytesIO()
        doc.save(temp_doc_io)
        temp_doc_io.seek(0)

        # Replace logo
        updated_doc_bytes = replace_image(
            temp_doc_io.read(), placeholder_logo_bytes, new_logo_bytes
        )

        # Re-open the modified document
        modified_doc = Document(io.BytesIO(updated_doc_bytes))

        # Apply text changes
        for template, localised in templates.items():
            replace_text_in_paragraphs(modified_doc.paragraphs, {template: localised})
            replace_text_in_tables(modified_doc.tables, {template: localised})
            replace_text_in_headers_footers(
                modified_doc.sections, {template: localised}
            )

        # Save modified document to a BytesIO object
        modified_doc_io = BytesIO()
        modified_doc.save(modified_doc_io)
        modified_doc_io.seek(0)
        modified_docs[doc_name] = modified_doc_io
    return modified_docs


### SESSION STATE ###

if "step" not in st.session_state:
    st.session_state.step = "pack"

if "zip_contents" not in st.session_state:
    st.session_state.zip_contents = None

if "placeholder_logo_name" not in st.session_state:
    st.session_state.placeholder_logo_name = None
if "placeholder_logo" not in st.session_state:
    st.session_state.placeholder_logo = None

if "uploaded_logo_status" not in st.session_state:
    st.session_state.uploaded_logo_status = False
if "uploaded_logo_name" not in st.session_state:
    st.session_state.uploaded_logo_name = None
if "uploaded_logo" not in st.session_state:
    st.session_state.uploaded_logo = None


if "docx_files" not in st.session_state:
    st.session_state.docx_files = {}

if "localisation_df_ready" not in st.session_state:
    st.session_state.localisation_df_ready = None
if "localisation_df_final" not in st.session_state:
    st.session_state.localisation_df_final = None


### MAIN / HEADING ###
st.title("🌐 LOCALE - Site")
col_caption, col_reset = st.columns([4, 1])  # Adjust column ratios as needed
with col_caption:
    st.caption(
        "LOcalisation of the Clinical triAls template fiLes with site-specific valuEs and site logo"
    )
with col_reset:
    if st.button("🔄 Reset tool"):
        st.session_state.step = "pack"
        st.session_state.zip_contents = None
        st.session_state.placeholder_logo_name = None
        st.session_state.placeholder_logo = None
        st.session_state.uploaded_logo_status = False
        st.session_state.uploaded_logo_name = None
        st.session_state.uploaded_logo = None
        st.session_state.docx_files = {}
        st.session_state.localisation_df_ready = None
        st.session_state.localisation_df_final = None
        st.rerun()

bosh = st.empty()

### MAIN / LOCALISATION PACK ###

if st.session_state.step == "pack":
    pack = bosh.container(border=True)
    with pack:
        st.markdown(
            """
Before starting the localisation, make sure that you have the following:

1. The localisation zip package provided by the trial coordination team 
2. Your site’s logo

You will be following these steps:

1. Uploading the zip package provided by the coordination team
2. Uploading the "localised.csv" file if you have used LOCALE for this trial before (optional)
3. Using the data editor to fill the missing or incorrect localisation values
4. Replacing the placeholder logo in the documents with your site’s logo
5. Finalising the data entry and downloading the localised documents

If you want more information before using, you can click [here](https://github.com/erdemdemir/locale/wiki/Step-by-step-instructions-for-the-Site-tool) to view the  step-by-step instructions.
        """
        )
        uploaded_package = st.file_uploader(
            "Upload the zip package provided by the coordination team", type="zip"
        )

        if uploaded_package is not None and st.session_state.zip_contents is None:
            # If a file is uploaded, read its contents and store in session state
            st.session_state.zip_contents = uploaded_package.getvalue()

        if st.session_state.zip_contents is not None:
            # Open the ZIP file from the session state
            with zipfile.ZipFile(
                io.BytesIO(st.session_state.zip_contents), "r"
            ) as package_zip:
                docx_files, placeholder_logo, localisation_csv = (
                    extract_package_contents(package_zip)
                )

                if not all([docx_files, placeholder_logo, localisation_csv]):
                    st.error(
                        "One or more required files are missing from the zip package. Please check the contents and try again."
                    )
                    st.stop()

                # Store the extracted image files in session state
                st.session_state.placeholder_logo_name = placeholder_logo
                with package_zip.open(placeholder_logo) as img_file:
                    st.session_state.placeholder_logo = img_file.read()

                # Read the localisation CSV file for further checks and processing
                with package_zip.open(localisation_csv) as csv_file:
                    localisation_df = pd.read_csv(csv_file)

                    # Ensure 'Localisation' column is present, if not, add it with empty values
                    if "localisation" not in localisation_df.columns:
                        localisation_df["localisation"] = ""

                    # Define expected columns in 'localisation.csv'
                    expected_columns = [
                        "document_type",
                        "document_name",
                        "variable",
                        "localisation",
                    ]
                    # Check if all expected columns are present, if not, display an error and stop
                    if not all(
                        col in localisation_df.columns for col in expected_columns
                    ):
                        st.error(
                            f"Missing expected columns in 'localisation.csv'. Expected: {', '.join(expected_columns)}"
                        )
                        st.stop()

                    # Group the DataFrame by 'Document' and convert it to a dictionary
                    localisation_dict = (
                        localisation_df.groupby("document_name")["variable"]
                        .apply(list)
                        .to_dict()
                    )

                # Check if document names in CSV match the ones in the zip
                docx_in_csv = set(localisation_dict.keys())
                docx_in_zip = set(docx_files)

                if docx_in_csv != docx_in_zip:
                    # Identify missing documents
                    missing_in_zip = docx_in_csv - docx_in_zip
                    missing_in_csv = docx_in_zip - docx_in_csv
                    st.error("Document lists do not match.")
                    if missing_in_zip:
                        st.error(f"Documents in CSV not in zip: {missing_in_zip}")
                    if missing_in_csv:
                        st.error(f"Documents in zip not in CSV: {missing_in_csv}")
                    st.stop()
                else:
                    st.success("The zip package looks valid 🎉")
                    data = []

                    # Create a dictionary mapping document names to their document types
                    document_type_dict = localisation_df.set_index("document_name")[
                        "document_type"
                    ].to_dict()

                    # Process each .docx file
                    for docx in docx_files:
                        docx_path = next(
                            f for f in package_zip.namelist() if f.endswith(docx)
                        )
                        with package_zip.open(docx_path) as file:
                            doc = Document(io.BytesIO(file.read()))
                            # Store the document in session state
                            st.session_state.docx_files[docx] = doc
                            # Iterate through paragraphs in the document
                            for para in doc.paragraphs:
                                for var in localisation_dict.get(docx, []):
                                    # If a variable is found in the paragraph, add it to the data list
                                    if var in para.text:
                                        data.append(
                                            {
                                                "document_type": document_type_dict.get(
                                                    docx, ""
                                                ),
                                                "document_name": docx,
                                                "related_paragraph": para.text,
                                                "variable": var,
                                                "localisation": "",
                                            }
                                        )

                    # Create a DataFrame from the extracted data
                    localisation_df_ready = pd.DataFrame(data)

                    if st.button("Click here to continue to the next step ➡️"):
                        st.session_state.localisation_df_ready = localisation_df_ready
                        st.session_state.step = "previous"
                        bosh.empty()
                        sleep(0.5)


### MAIN / PREVIOUS ATTEMPT ###

if st.session_state.step == "previous":
    # bosh.empty()
    previous = bosh.container(border=True)
    with previous:
        st.markdown(
            "If you have previously used LOCALE for this trial, such as for a past amendment, you can upload the 'localisation.csv' file from previous attempt here. Previously entered values will be imported to populate the matching fields, which in turn will save you some typing. This 'localisation.csv' file can be found in the zip package you downloaded at the end of that previous localisation alongside with the localised docx files."
        )

        # Load the localisation_df ready to update with previous values
        localisation_df_ready = st.session_state.localisation_df_ready

        # Ask user if they have a previous localisation.csv
        previous_attempt = st.radio(
            "Did you use LOCALE for this trial before?", ("Yes", "No")
        )

        if previous_attempt == "Yes":
            uploaded_file = st.file_uploader(
                "Please upload the 'localised.csv' file to import your previous values",
                type=["csv"],
            )
            if uploaded_file is not None:
                # Load the previous CSV
                previous_csv = pd.read_csv(uploaded_file)

                # Ensure 'localisation' column exists in previous_csv
                if "localisation" not in previous_csv.columns:
                    st.error(
                        "The uploaded CSV does not contain a 'localisation' column."
                    )
                else:
                    # Update the current DataFrame with 'localisation' values from previous_csv
                    for index, row in localisation_df_ready.iterrows():
                        # Match document type and variable
                        previous_data = previous_csv[
                            (previous_csv["document_type"] == row["document_type"])
                            & (previous_csv["variable"] == row["variable"])
                        ]

                        # Check if there are matching rows and update 'localisation'
                        if (
                            not previous_data.empty
                            and "localisation" in previous_data.columns
                        ):
                            # Copying the 'localisation' value
                            localisation_df_ready.at[index, "localisation"] = (
                                previous_data.iloc[0]["localisation"]
                            )

                    st.success(
                        "Previous localisation values have been successfully imported."
                    )
            if st.button("Click here to continue to data editor ➡️"):
                st.session_state.localisation_df_ready = localisation_df_ready
                st.session_state.step = "editor"
                bosh.empty()
                sleep(0.5)
        else:
            if st.button(
                "Click here to continue to data editor without importing previous values ➡️"
            ):
                st.session_state.localisation_df_ready = localisation_df_ready
                st.session_state.step = "editor"
                bosh.empty()
                sleep(0.5)

### MAIN / DATA EDITOR ###

if st.session_state.step == "editor":
    # bosh.empty()
    editor = bosh.container(border=True)
    with editor:
        st.markdown(
            """
You can use the data editor below to fill the missing values or correct mistakes. To be able to see the whole table and make edits, you must hover your mouse on the table and then click the rightmost fullscreen icon on top of the data editor.

The data editor has the following columns:

1. **Document type:** The type of the document
2. **Document name:** The name of the file
3. **Related paragraph:** The paragraph where the variable is found in the document to help you identify the variable
4. **Variable:** The variable to be localised
5. **Localisation:** The value to replace the variable with

If you made significant edits, I would strongly suggest to download the csv file before finalising data entry and move to the next stage. 
        """
        )

        # New dataframe with all localisation information
        localisation_df_final = st.data_editor(
            data=st.session_state.localisation_df_ready,
            column_config={
                "document_type": st.column_config.Column(
                    "Document type",
                    help="The type of the document",
                    width="medium",
                    disabled=True,
                ),
                "document_name": st.column_config.Column(
                    "Document name",
                    help="The name of the file",
                    width="medium",
                    disabled=True,
                ),
                "related_paragraph": st.column_config.Column(
                    "Related paragraph",
                    help="The paragraph where the variable is found in the document to help you identify the variable",
                    width="large",
                    disabled=True,
                ),
                "variable": st.column_config.Column(
                    "Variable",
                    help="The variable to be localised",
                    width="medium",
                    disabled=True,
                ),
                "localisation": st.column_config.Column(
                    "Localisation",
                    help="The value to replace the variable with",
                    width="large",
                ),
            },
        )

        if st.button("Click to finalise the data entry and move to logo replacement ➡️"):
            st.session_state.localisation_df_final = localisation_df_final
            st.session_state.step = "logo"
            bosh.empty()
            sleep(0.5)

### MAIN / LOGO REPLACEMENT ###

if st.session_state.step == "logo":
    # bosh.empty()
    logo = bosh.container(border=True)
    with logo:
        placeholder_logo_name = st.session_state.placeholder_logo_name
        placeholder_logo = st.session_state.placeholder_logo
        width, height = Image.open(io.BytesIO(placeholder_logo)).size

        st.markdown(
            f"You can update the placeholder logo with your site’s logo. To start, upload your logo. The expected width and height of the site logo are :red[{width}] and :red[{height}] pixels, respectively."
        )

        file_name = placeholder_logo_name.split("/")[
            -1
        ]  # Extract the file name for display

        with st.expander("👉 Click here if you like to view the placeholder logo"):
            st.image(st.session_state.placeholder_logo, caption=file_name)
        uploaded_logo = st.file_uploader(
            f"Upload a new logo to replace placeholde logo",
            type=["png", "jpg", "jpeg"],
            key=f"upload",
        )

        if uploaded_logo is not None:

            uploaded_logo_bytes = uploaded_logo.read()
            st.session_state.uploaded_logo = uploaded_logo_bytes

            st.session_state.uploaded_logo_name = uploaded_logo.name
            st.session_state.uploaded_logo_status = True

            # Display the placeholder logo and the uploaded logo side by side
            cols = st.columns(2)
            with cols[0]:
                st.image(st.session_state.placeholder_logo, caption=file_name)
            with cols[1]:
                if uploaded_logo is not None:
                    st.image(uploaded_logo_bytes, caption=uploaded_logo.name)
                else:
                    st.write("No replacement uploaded.")

        # Display a submit button for final confirmation
        uploaded_logo_status = st.session_state.uploaded_logo_status
        if uploaded_logo_status is True:
            if st.button(
                "I agree, submit this and continue creating the localised documents ➡️"
            ):
                st.write("Logos have been uploaded successfully.")
                st.session_state.step = "finalise"
                bosh.empty()
                sleep(0.5)


if st.session_state.step == "finalise":
    final = bosh.container(border=True)
    with final:
        # Load the localisation_df_final to extract the changes
        localisation_df_final = st.session_state.localisation_df_final

        # Create JSON structure from DataFrame
        changes_json = {}
        for _, row in localisation_df_final.iterrows():
            document_name = row["document_name"]
            variable = row["variable"]
            localisation = row["localisation"]
            if document_name not in changes_json:
                changes_json[document_name] = {}
            changes_json[document_name][variable] = localisation

        # Document files to apply changes to
        docx_files = st.session_state.docx_files

        st.write("The following documents will be modified:")
        for doc_name in docx_files:
            st.write(f"- {doc_name}")

        if st.button("Apply changes and create the localised documents for download"):
            modified_docs = apply_changes_to_document(
                changes_json,
                docx_files,
                st.session_state.placeholder_logo,
                st.session_state.uploaded_logo,
            )

            # Create a ZIP file with all modified documents and the localisation CSV
            zip_buffer = io.BytesIO()
            with zipfile.ZipFile(zip_buffer, "w") as zip_file:
                for doc_name, doc_io in modified_docs.items():
                    zip_file.writestr(f"localised_{doc_name}", doc_io.getvalue())
                localisation_csv = localisation_df_final.to_csv(index=False)
                zip_file.writestr("localisation.csv", localisation_csv)
            zip_buffer.seek(0)

            st.success(
                "The localised documents have been successfully created. You can now download the ZIP package using the button below 🎉"
            )

            # Provide a single download button for the ZIP file
            st.download_button(
                label="Download localised documents as a zip package ⬇️",
                data=zip_buffer,
                file_name="localised_documents.zip",
                mime="application/zip",
            )

# type: ignore
